# -*- coding: utf-8 -*-
"""
Recorder de ações com relatório visual (Markdown + HTML + opcional DOCX)
- Marca o ponto clicado na imagem (círculo)
- Registra título da janela ativa e (no Windows) nome do app
- Gera:
    recording_YYYYMMDD_HHMMSS/
      ├─ images/                (originais)
      ├─ images_marked/         (com círculo no clique)
      ├─ steps.md               (markdown)
      ├─ report.html            (relatório visual)
      └─ Relatorio_Acoes.docx   (se python-docx disponível)

Atalhos:
  ESC = encerrar
  F9  = screenshot manual (sem clique)

Segurança/Privacidade: não execute com dados sensíveis visíveis.
"""

import os
import time
import platform
from datetime import datetime
from pathlib import Path
import threading
from dataclasses import dataclass, asdict
from typing import Optional, List, Tuple

# deps de captura
import pyautogui
from pynput import mouse, keyboard
from PIL import Image, ImageDraw

# deps janela ativa
import psutil
try:
    import pygetwindow as gw
except Exception:
    gw = None

is_windows = platform.system().lower().startswith("win")
if is_windows:
    try:
        import win32gui  # type: ignore
        import win32process  # type: ignore
    except Exception:
        win32gui = None     # type: ignore
        win32process = None # type: ignore

# opcional DOCX
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# =========================
# CONFIGURAÇÕES
# =========================
OUTPUT_ROOT = Path.cwd()  # pode trocar
CROP_RADIUS = 0           # 0 = tela cheia; ex: 220 para recorte ao redor do clique
ANNOTATE_CLICK = True
CIRCLE_RADIUS = 18
CIRCLE_WIDTH = 5
CIRCLE_COLOR = (220, 40, 40)   # vermelho
SHADOW_COLOR = (0, 0, 0, 90)   # alpha para sombra

HTML_THEME_DARK = True         # tema escuro do relatório

# =========================
# ESTRUTURA DE SAÍDA
# =========================
RUN_ID = datetime.now().strftime("%Y%m%d_%H%M%S")
BASE_OUT = OUTPUT_ROOT / f"recording_{RUN_ID}"
IMAGES_DIR = BASE_OUT / "images"
IMAGES_MARKED_DIR = BASE_OUT / "images_marked"
MD_FILE = BASE_OUT / "steps.md"
HTML_FILE = BASE_OUT / "report.html"
DOCX_FILE = BASE_OUT / "Relatorio_Acoes.docx"

for d in (IMAGES_DIR, IMAGES_MARKED_DIR):
    d.mkdir(parents=True, exist_ok=True)

# =========================
# ESTADO
# =========================
@dataclass
class Step:
    idx: int
    ts: str
    action: str         # "click left", "scroll", "manual"
    x: Optional[int]
    y: Optional[int]
    window_title: str
    app_name: str
    img_rel: str        # caminho relativo imagem original
    img_mark_rel: str   # imagem anotada (círculo), se houver

steps: List[Step] = []
running = True
lock = threading.Lock()
step_counter = 0

# =========================
# HELPERS
# =========================
def now_human() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def active_window_info() -> Tuple[str, str]:
    """Retorna título da janela ativa e nome do processo (quando possível)."""
    title, app = "", ""
    # pygetwindow (multi-plataforma)
    if gw is not None:
        try:
            win = gw.getActiveWindow()
            if win:
                title = (win.title or "").strip()
                if is_windows and getattr(win, "_hWnd", None) and win32gui and win32process:
                    try:
                        hwnd = win._hWnd
                        _, pid = win32process.GetWindowThreadProcessId(hwnd)
                        if pid:
                            app = psutil.Process(pid).name()
                    except Exception:
                        pass
        except Exception:
            pass

    # fallback Windows puro
    if not title and is_windows and win32gui:
        try:
            hwnd = win32gui.GetForegroundWindow()
            if hwnd:
                title = (win32gui.GetWindowText(hwnd) or "").strip()
                if win32process:
                    try:
                        _, pid = win32process.GetWindowThreadProcessId(hwnd)
                        if pid:
                            app = psutil.Process(pid).name()
                    except Exception:
                        pass
        except Exception:
            pass

    return title, app

def screenshot_path(prefix: str = "step") -> Path:
    return IMAGES_DIR / f"{prefix}_{int(time.time()*1000)}.png"

def annotate_click_on_image(src: Path, dst: Path, click_xy: Optional[Tuple[int,int]]):
    im = Image.open(src).convert("RGBA")
    if click_xy is not None:
        x, y = click_xy
        draw = ImageDraw.Draw(im, "RGBA")
        # sombra
        draw.ellipse(
            (x - CIRCLE_RADIUS - 2, y - CIRCLE_RADIUS - 2,
             x + CIRCLE_RADIUS + 2, y + CIRCLE_RADIUS + 2),
            outline=SHADOW_COLOR, width=CIRCLE_WIDTH+2
        )
        # círculo
        draw.ellipse(
            (x - CIRCLE_RADIUS, y - CIRCLE_RADIUS,
             x + CIRCLE_RADIUS, y + CIRCLE_RADIUS),
            outline=CIRCLE_COLOR, width=CIRCLE_WIDTH
        )
    im.convert("RGB").save(dst)

def capture(region=None) -> Path:
    p = screenshot_path()
    pyautogui.screenshot(region=region).save(p)
    return p

def capture_with_optional_crop(x: Optional[int], y: Optional[int]) -> Path:
    if CROP_RADIUS and (x is not None and y is not None):
        left = max(0, x - CROP_RADIUS)
        top  = max(0, y - CROP_RADIUS)
        width = CROP_RADIUS * 2
        height = CROP_RADIUS * 2
        return capture(region=(left, top, width, height))
    return capture()

def record(action: str, x: Optional[int], y: Optional[int], img_path: Path):
    global step_counter
    with lock:
        step_counter += 1
        title, app = active_window_info()
        marked_rel = ""

        # gerar imagem marcada
        if ANNOTATE_CLICK and action.startswith("click"):
            marked = IMAGES_MARKED_DIR / img_path.name
            try:
                annotate_click_on_image(img_path, marked, (x, y))
                marked_rel = f"images_marked/{marked.name}"
            except Exception:
                marked_rel = ""

        step = Step(
            idx=step_counter,
            ts=now_human(),
            action=action,
            x=x, y=y,
            window_title=title,
            app_name=app,
            img_rel=f"images/{img_path.name}",
            img_mark_rel=marked_rel or f"images/{img_path.name}"
        )
        steps.append(step)
        print(f"[{step.ts}] {action} ({x},{y}) -> {img_path.name} | {title} ({app})")

# =========================
# LISTENERS
# =========================
def on_click(x, y, button, pressed):
    if not pressed:
        return
    btn = str(button).replace("Button.", "")
    img = capture_with_optional_crop(x, y)
    record(f"click {btn}", x, y, img)

def on_scroll(x, y, dx, dy):
    img = capture_with_optional_crop(x, y)
    record(f"scroll dx={dx} dy={dy}", x, y, img)

def on_press(key):
    global running
    try:
        if key == keyboard.Key.esc:
            running = False
            return False
        if key == keyboard.Key.f9:
            img = capture()
            record("manual screenshot", None, None, img)
    except Exception as e:
        print("on_press error:", e)

# =========================
# RELATÓRIOS
# =========================
def write_markdown():
    lines = []
    lines.append(f"# Registro de Ações — {RUN_ID}\n")
    lines.append("> **Dica**: substitua os campos *Objetivo/Observação* após a captura.\n\n")
    for s in steps:
        lines.append(f"## Passo {s.idx}\n")
        lines.append(f"**{s.ts}**\n\n")
        if s.window_title or s.app_name:
            w = f"**Janela:** {s.window_title}" if s.window_title else ""
            a = f" (app: `{s.app_name}`)" if s.app_name else ""
            lines.append(f"- {w}{a}\n")
        if s.action.startswith("click"):
            lines.append(f"- **Ação:** {s.action} — posição ({s.x}, {s.y})\n")
        else:
            lines.append(f"- **Ação:** {s.action}\n")
        lines.append(f"- **Objetivo/Observação:** _preencha aqui_\n\n")
        lines.append(f"![screenshot]({s.img_mark_rel})\n\n")
    MD_FILE.write_text("\n".join(lines), encoding="utf-8")

def write_html():
    css_dark = """
    :root { --bg:#111; --fg:#eaeaea; --muted:#9aa0a6; --card:#1b1b1b; --accent:#7aa2ff; --border:#2a2a2a; }
    """
    css_light = """
    :root { --bg:#f7f7f9; --fg:#1a1a1a; --muted:#616161; --card:#ffffff; --accent:#2a66ff; --border:#e6e6e6; }
    """
    css = css_dark if HTML_THEME_DARK else css_light
    head = f"""<!doctype html>
<html lang="pt-br"><head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>Relatório de Ações — {RUN_ID}</title>
<style>
{css}
*{{box-sizing:border-box}}
body{{margin:0;background:var(--bg);color:var(--fg);font:16px/1.5 system-ui,Segoe UI,Roboto,Arial}}
.container{{max-width:1120px;margin:40px auto;padding:0 20px}}
h1{{font-size:28px;margin:0 0 8px}}
.subtitle{{color:var(--muted);margin-bottom:24px}}
.index{{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:12px;margin:16px 0 28px}}
.badge{{display:inline-block;padding:2px 8px;background:var(--accent);color:#fff;border-radius:999px;font-size:12px;margin-left:8px}}
.card{{background:var(--card);border:1px solid var(--border);border-radius:16px;padding:16px;margin:14px 0;box-shadow:0 2px 8px rgba(0,0,0,.15)}}
.card h2{{margin:0 0 6px;font-size:20px}}
.meta{{color:var(--muted);font-size:13px;margin-bottom:10px}}
.imgwrap{{margin-top:10px;border-radius:12px;overflow:hidden;border:1px solid var(--border)}}
.imgwrap img{{width:100%;display:block}}
.hr{{height:1px;background:var(--border);margin:28px 0}}
.toplink{{font-size:13px;color:var(--accent);text-decoration:none}}
.kv{{display:grid;grid-template-columns:120px 1fr;gap:8px}}
.kv div:first-child{{color:var(--muted)}}
.footer{{color:var(--muted);font-size:12px;margin:32px 0}}
a{ color:var(--accent); }
</style>
</head><body><div class="container">
<h1>Relatório de Ações <span class="badge">{RUN_ID}</span></h1>
<div class="subtitle">Gerado automaticamente. Pressione ESC para encerrar; F9 captura manual.</div>
<h3>Índice</h3>
<div class="index">
"""
    idx = []
    for s in steps:
        idx.append(f'<a class="toplink" href="#step-{s.idx}">Passo {s.idx} — {s.ts}</a>')
    head += "\n".join(idx) + "</div><div class='hr'></div>\n"

    body = []
    for s in steps:
        body.append(f'<div class="card" id="step-{s.idx}">')
        body.append(f"<h2>Passo {s.idx}</h2>")
        body.append(f'<div class="meta">{s.ts}</div>')
        body.append('<div class="kv">')
        if s.window_title:
            body.append(f"<div>Janela</div><div>{s.window_title}</div>")
        if s.app_name:
            body.append(f"<div>Aplicativo</div><div><code>{s.app_name}</code></div>")
        if s.action.startswith("click"):
            body.append(f"<div>Ação</div><div>click — posição ({s.x}, {s.y})</div>")
        else:
            body.append(f"<div>Ação</div><div>{s.action}</div>")
        body.append(f"<div>Observação</div><div><em>preencha aqui</em></div>")
        body.append("</div>")
        body.append(f'<div class="imgwrap"><img src="{s.img_mark_rel}" alt="Passo {s.idx}"/></div>')
        body.append('<div style="margin-top:8px"><a class="toplink" href="#topo">↑ voltar ao topo</a></div>')
        body.append("</div>")

    foot = f"""
<div class="footer">Relatório gerado em {now_human()} — {len(steps)} passo(s).</div>
</div></body></html>
"""
    HTML_FILE.write_text(head + "\n".join(body) + foot, encoding="utf-8")

def write_docx_if_possible():
    if not HAS_DOCX:
        print("python-docx não instalado — pulando DOCX (pip install python-docx)")
        return
    doc = Document()

    # capa simples
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Relatório de Ações")
    run.bold = True
    run.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"ID: {RUN_ID} — Gerado em {now_human()}")
    run.font.size = Pt(12)

    doc.add_page_break()

    for s in steps:
        h = doc.add_heading(f"Passo {s.idx}", level=1)
        pmeta = doc.add_paragraph(f"{s.ts}")
        if s.window_title:
            doc.add_paragraph(f"Janela: {s.window_title}")
        if s.app_name:
            doc.add_paragraph(f"Aplicativo: {s.app_name}")
        if s.action.startswith("click"):
            doc.add_paragraph(f"Ação: {s.action} — posição ({s.x}, {s.y})")
        else:
            doc.add_paragraph(f"Ação: {s.action}")
        doc.add_paragraph("Observação: ______________________________")

        imgp = BASE_OUT / s.img_mark_rel
        try:
            doc.add_picture(str(imgp), width=Inches(6.2))
        except Exception:
            pass
        doc.add_page_break()

    doc.save(DOCX_FILE)
    print(f"📄 DOCX gerado: {DOCX_FILE}")

# =========================
# LOOP PRINCIPAL
# =========================
def start():
    # markdown cabeçalho provisório (será reescrito ao final)
    MD_FILE.write_text(f"# Registro de Ações — {RUN_ID}\n\n", encoding="utf-8")

    m_listener = mouse.Listener(on_click=on_click, on_scroll=on_scroll)
    k_listener = keyboard.Listener(on_press=on_press)
    m_listener.start(); k_listener.start()
    print("Iniciado. ESC para encerrar, F9 para screenshot manual.")
    print(f"Saída: {BASE_OUT}")

    try:
        while running:
            time.sleep(0.2)
    finally:
        m_listener.stop(); k_listener.stop()

    # gerar relatórios
    write_markdown()
    write_html()
    write_docx_if_possible()
    print("\n✅ Finalizado!")
    print(f"- Markdown: {MD_FILE}")
    print(f"- HTML:     {HTML_FILE}")
    if HAS_DOCX:
        print(f"- DOCX:     {DOCX_FILE}")

if __name__ == "__main__":
    start()
